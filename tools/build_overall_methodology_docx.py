from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "results" / "Overall_Methodology_AUX_ESN_XGBoost.docx"

BLUE = "1F5A7A"
DARK = "17324D"
MID = "587084"
LIGHT_BLUE = "EAF3F8"
PALE_TEAL = "EAF7F4"
PALE_GOLD = "FFF6DF"
PALE_RED = "FDEEEF"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
BLACK = "000000"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total = int(sum(widths_inches) * 1440)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for old in list(grid):
        grid.remove(old)
    for width in widths_inches:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_inches):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run(run, size=11, bold=False, color=BLACK, italic=False, font="Calibri"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_body(doc, text, bold_lead=None, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, bold=True, color=DARK)
        r = p.add_run(text[len(bold_lead):])
        set_run(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    set_run(p.add_run(text), size=10.5)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.16)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), BLUE)
    left.set(qn("w:space"), "6")
    borders.append(left)
    p_pr.append(borders)
    set_run(p.add_run(label + "  "), bold=True, color=BLUE)
    set_run(p.add_run(text), color=DARK)
    return p


def add_equation(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    set_run(p.add_run(text), size=12, italic=True, color=DARK, font="Cambria Math")
    return p


def add_table(doc, headers, rows, widths, header_fill=BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run(p.add_run(header), size=10, bold=True, color=WHITE)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_run(p.add_run(str(value)), size=9.5, color=DARK)
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_step(doc, number, title, purpose, details, equation=None, note=None):
    h = doc.add_heading(f"Step {number}. {title}", level=2)
    h.paragraph_format.keep_with_next = True
    add_body(doc, purpose, bold_lead="Purpose: ")
    for detail in details:
        add_bullet(doc, detail)
    if equation:
        add_equation(doc, equation)
    if note:
        add_callout(doc, "Important", note, PALE_GOLD)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

# Compact reference guide preset tokens.
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25
for name, size, color, before, after in (
    ("Title", 30, DARK, 0, 8),
    ("Subtitle", 14, MID, 0, 8),
    ("Heading 1", 17, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 11.5, DARK, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = name != "Subtitle"
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

# List rhythm.
for name in ("List Bullet", "List Bullet 2", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(10.5)
    styles[name].paragraph_format.space_after = Pt(4)
    styles[name].paragraph_format.line_spacing = 1.2

# Running header.
header = section.header
p = header.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(p.add_run("AUX ESN–XGBoost Methodology"), size=9, bold=True, color=MID)

# Footer with page field.
footer = section.footer
p = footer.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(p.add_run("Page "), size=9, color=MID)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
p._p.append(fld)

# Cover.
doc.add_paragraph().paragraph_format.space_after = Pt(70)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
set_run(p.add_run("METHODOLOGY GUIDE"), size=11, bold=True, color=BLUE)
p = doc.add_paragraph(style="Title")
set_run(p.add_run("From Thermal Contact to Predicted Effusivity"), size=30, bold=True, color=DARK)
p = doc.add_paragraph(style="Subtitle")
set_run(p.add_run("Complete AUX → AUX-BB ESN–XGBoost workflow"), size=15, color=MID)
add_body(doc, "A beginner-friendly reference covering data preparation, scaling, reservoir dynamics, feature construction, Optuna optimization, cross-validation, performance metrics, and final deployment.", keep=True)
doc.add_paragraph().paragraph_format.space_after = Pt(24)
add_callout(doc, "Core question", "Can the first five seconds of the post-contact thermal response predict the effusivity of a material that was not used for training?", PALE_TEAL)
add_table(doc, ["Dataset component", "Count"], [
    ("Materials", "14"),
    ("Operating temperatures", "4: 30°C, 40°C, 50°C and 60°C"),
    ("Repetitions per material and temperature", "6"),
    ("Total trials", "14 × 4 × 6 = 336"),
], [2.8, 3.7])
add_body(doc, "Prepared for the current baseline-centered multi-temperature modeling workflow.", italic=True)

doc.add_page_break()
doc.add_heading("1. Methodology at a Glance", level=1)
add_callout(doc, "One-sentence summary", "The sensors record a thermal transient; baseline-centered and standardized channels drive an ESN; trajectory summaries become XGBoost predictors; leave-one-material-out testing measures generalization to unseen materials.")

flow_rows = [
    ("1", "336 raw thermal trials", "Primary and Secondary time series"),
    ("2", "Assign material properties", "Compute effusivity e = √(kρcₚ)"),
    ("3", "Detect and align contact", "Set contact to t = 0"),
    ("4", "Extract analysis window", "Retain 0–5 s after contact"),
    ("5", "Smooth and baseline-center", "Express change relative to first retained sample"),
    ("6", "Construct five ESN channels", "Two signals, difference and two derivatives"),
    ("7", "Create a LOMO fold", "312 training trials; 24 held-out trials"),
    ("8", "Fit training-only scaler", "Standardize five channels without test leakage"),
    ("9", "Run three reservoirs", "Seeds 42, 43 and 44"),
    ("10", "Summarize trajectories", "Nine summaries per reservoir unit"),
    ("11", "Train three XGBoost models", "Predict the held-out material"),
    ("12", "Average seed predictions", "One prediction per test trial"),
    ("13", "Pool all LOMO predictions", "336 out-of-fold predictions"),
    ("14", "Score the candidate", "NRMSE, R², fold stability and objective J"),
    ("15", "Repeat Optuna search", "Evaluate 100 joint ESN–XGBoost candidates"),
    ("16", "Load best settings in AUX-BB", "Run fixed diagnostics and export results"),
]
add_table(doc, ["#", "Stage", "Output"], flow_rows, [0.45, 2.55, 3.5])

doc.add_heading("2. What the Model Predicts", level=1)
add_body(doc, "The model predicts thermal effusivity, a continuous material property describing how strongly a material exchanges heat with the sensor during contact.")
add_equation(doc, "e = √(k ρ cₚ)")
add_body(doc, "One complete 0–5 s trial produces one predicted effusivity value. This is regression—not timestep-by-timestep classification.")

doc.add_page_break()
doc.add_heading("3. Data Preparation", level=1)
add_step(doc, 1, "Load and identify every trial", "Create the complete multi-temperature dataset.", [
    "Read valid CSV files from the 30C, 40C, 50C and 60C folders.",
    "Construct a unique trial ID containing temperature, material and repetition.",
    "Use one authoritative material-property table instead of repeated processed-file values.",
    "Calculate one true effusivity for every material.",
])
add_step(doc, 2, "Detect contact", "Establish a common physical starting point.", [
    "Smooth the Primary sensor signal.",
    "Calculate its derivative and locate the strongest downward transition.",
    "Identify the beginning of that transition as the estimated contact time.",
    "No k or effusivity value is used to locate contact.",
], equation="t_aligned = t_original − t_contact")
add_step(doc, 3, "Extract the analysis window", "Focus the model on the post-contact thermal transient.", [
    "Set the aligned contact point to t = 0.",
    "Retain samples from 0 through 5 seconds after contact.",
    "Reject a trial if too few valid post-contact samples remain.",
], note="All later preprocessing, ESN dynamics and feature summaries refer to this same 0–5 s window.")

doc.add_heading("4. Sensor Preprocessing", level=1)
add_step(doc, 4, "Smooth Primary and Secondary", "Reduce small measurement noise without discarding the transient shape.", [
    "Apply a Savitzky–Golay filter separately to Primary and Secondary.",
    "Preserve the timing and overall curvature of the response.",
])
add_step(doc, 5, "Baseline-center each trial", "Remove the absolute starting offset and focus on change after contact.", [
    "Subtract the first retained smoothed value from every Primary sample.",
    "Repeat independently for Secondary.",
    "Each trial uses its own baseline; no target information is used.",
], equation="PΔ(t) = Pₛ(t) − Pₛ(t₀)     and     SΔ(t) = Sₛ(t) − Sₛ(t₀)")
add_callout(doc, "Presentation versus model input", "Multiplying changes by 10⁶ makes plots readable, but this multiplier is not used by the model. The following StandardScaler would cancel a constant multiplier.", PALE_GOLD)

doc.add_page_break()
doc.add_heading("5. Construct the Five ESN Input Channels", level=1)
add_table(doc, ["Channel", "Definition", "Interpretation"], [
    ("1", "PΔ(t)", "Baseline-centered Primary response"),
    ("2", "SΔ(t)", "Baseline-centered Secondary response"),
    ("3", "PΔ(t) − SΔ(t)", "Difference between sensor changes"),
    ("4", "dPΔ(t)/dt", "Local Primary rate of change"),
    ("5", "dSΔ(t)/dt", "Local Secondary rate of change"),
], [0.65, 1.75, 4.1])
add_body(doc, "At every timestep, the reservoir receives a five-number vector u(t). The derivative channels are local slopes, not one overall slope for the complete trial.")

doc.add_heading("6. Leakage-Safe Cross-Validation Split", level=1)
add_step(doc, 6, "Create one leave-one-material-out fold", "Test prediction on a material completely absent from training.", [
    "Hold out one material across all four temperatures and six repetitions.",
    "Train on the remaining 13 materials: 13 × 4 × 6 = 312 trials.",
    "Test on the held-out material: 1 × 4 × 6 = 24 trials.",
    "Repeat until all 14 materials have been held out once.",
])
add_callout(doc, "Concrete example", "When aluminum is held out, the scaler, ESN feature construction and XGBoost fitting use the other 13 materials. Aluminum contributes only test inputs and known targets used after prediction for scoring.", PALE_TEAL)

add_step(doc, 7, "Fit the five-channel StandardScaler", "Place channels on comparable numerical scales without using held-out information.", [
    "Stack all five-channel sequences from the 312 training trials.",
    "Calculate one training mean and one training standard deviation per channel.",
    "Transform both training and test channels using those training-derived values.",
    "Never refit the scaler on the held-out material.",
], equation="zⱼ(t) = [uⱼ(t) − μⱼ,train] / σⱼ,train")

doc.add_page_break()
doc.add_heading("7. Echo State Network Processing", level=1)
add_step(doc, 8, "Create the reservoir", "Transform five sensor channels into many nonlinear temporal state trajectories.", [
    "Input weights connect the five channels and a bias term to every reservoir unit.",
    "Recurrent weights connect previous reservoir states to the next state.",
    "The weights are randomly initialized and then fixed; they are not trained by XGBoost.",
    "Once the seed, inputs and hyperparameters are fixed, the trajectories are deterministic.",
])
add_equation(doc, "x̃(t) = tanh(Wᵢₙ[1; z(t)] + W x(t−1))")
add_equation(doc, "x(t) = (1−α)x(t−1) + αx̃(t)")
add_table(doc, ["ESN parameter", "Meaning"], [
    ("Reservoir size", "Number of internal state trajectories"),
    ("Leak rate α", "Speed of state updating versus memory of the previous state"),
    ("Input magnitude", "Strength of the five scaled inputs entering the reservoir"),
    ("Spectral radius", "Strength and persistence of recurrent state interactions"),
    ("Washout", "Initial recorded states discarded before summarization"),
], [1.65, 4.85])

add_step(doc, 9, "Run three reservoir seeds", "Reduce sensitivity to one random reservoir initialization.", [
    "Seed 42 generates one reservoir and feature table.",
    "Seeds 43 and 44 repeat the process with different fixed random weights.",
    "All three use the same optimized ESN hyperparameters and the same trial split.",
])
add_step(doc, 10, "Apply washout", "Reduce direct influence of the artificial initial zero state.", [
    "The reservoir processes every input sample.",
    "Only the first recorded states are excluded from trajectory summaries.",
    "Later retained states still contain recurrent information from earlier inputs.",
], note="With washout = 5, the first five recorded state rows are removed, not the first five input samples from reservoir processing.")

doc.add_page_break()
doc.add_heading("8. Convert Dynamics into XGBoost Predictors", level=1)
add_step(doc, 11, "Summarize each reservoir trajectory", "Convert a time series into a fixed-length feature row.", [
    "For every reservoir unit, calculate mean, standard deviation, range, net change and linear slope.",
    "Also calculate absolute area, time of maximum absolute activation, skewness and excess kurtosis.",
    "Add the known operating temperature, Temperature_C.",
])
add_equation(doc, "Predictor count = 9 × reservoir size + 1 temperature predictor")
add_table(doc, ["Reservoir size", "ESN summaries", "Total with temperature"], [
    ("20", "180", "181"),
    ("25", "225", "226"),
    ("69", "621", "622"),
], [1.8, 2.2, 2.5])
add_callout(doc, "Interpretation", "The representative state curves and heatmap are descriptive. XGBoost receives the numerical trajectory summaries—not the plotted image.")

doc.add_heading("9. XGBoost Regression", level=1)
add_step(doc, 12, "Train one XGBoost regressor per seed", "Learn the nonlinear mapping from ESN summaries to effusivity.", [
    "Train XGBoost 42 on the seed-42 training feature rows.",
    "Repeat independently for seeds 43 and 44.",
    "The held-out material's target values are never used for fitting.",
    "Transform the positive target with log1p during fitting and return predictions with expm1.",
    "Clip predictions to the target range observed in that fold's training data.",
])
add_step(doc, 13, "Average the three predictions", "Produce one robust prediction per held-out trial.", [
    "Each seed-specific regressor predicts the same 24 held-out trials.",
    "Average the three predictions trial by trial.",
], equation="ŷᵢ = (ŷᵢ,42 + ŷᵢ,43 + ŷᵢ,44) / 3")

doc.add_page_break()
doc.add_heading("10. Out-of-Fold Evaluation", level=1)
add_step(doc, 14, "Complete all 14 LOMO folds", "Ensure every material is evaluated as unseen exactly once.", [
    "Each fold produces 24 averaged test predictions.",
    "Fourteen folds produce 14 × 24 = 336 out-of-fold predictions.",
    "For one hyperparameter candidate, 14 folds × 3 seeds = 42 fitted XGBoost regressors.",
])
add_step(doc, 15, "Pool the predictions", "Calculate one overall unseen-material performance result.", [
    "Concatenate predictions from aluminum, bismuth, cement and all remaining held-out materials.",
    "Calculate pooled metrics using all 336 true and predicted values.",
    "Retain material-fold errors to diagnose unstable or difficult materials.",
])

doc.add_heading("11. Performance Metrics", level=1)
add_table(doc, ["Metric", "Meaning", "Preferred direction"], [
    ("MAE", "Average absolute prediction error", "Lower"),
    ("RMSE", "Error emphasizing larger mistakes", "Lower"),
    ("NRMSE range", "RMSE divided by the full target range", "Lower; 0 is perfect"),
    ("R²", "Improvement over predicting the pooled target mean", "Higher; 1 is perfect"),
    ("Median APE", "Median absolute percentage error", "Lower"),
    ("Fold-error SD", "Variation in normalized error among held-out materials", "Lower"),
], [1.25, 3.85, 1.4])
add_equation(doc, "R² = 1 − Σ(yᵢ−ŷᵢ)² / Σ(yᵢ−ȳ)²")
add_callout(doc, "Why per-material R² is invalid", "Every trial of one material has the same true effusivity, so the within-material target variance is zero. Report pooled LOMO R² and material-level error metrics instead.", PALE_GOLD)

doc.add_page_break()
doc.add_heading("12. Joint Hyperparameter Optimization in AUX", level=1)
add_step(doc, 16, "Let Optuna nominate one joint candidate", "Search ESN and XGBoost settings together.", [
    "ESN search variables: reservoir size, leak rate, input magnitude, spectral radius and washout.",
    "XGBoost variables: tree count, depth, learning rate, minimum child weight, row and column subsampling, and L1/L2 regularization.",
    "The multivariate TPE sampler learns which parameter regions tend to produce better objective values.",
])
add_step(doc, 17, "Evaluate the candidate through complete LOMO", "Judge every proposed configuration using the same unseen-material procedure.", [
    "Run all 14 folds with the candidate.",
    "Generate three seed-specific models per fold.",
    "Pool all 336 predictions and measure accuracy and fold stability.",
])
add_equation(doc, "J = pooled NRMSE + 0.25 × SD(material-fold normalized RMSE)")
add_step(doc, 18, "Repeat for 100 Optuna trials", "Improve the search efficiently without enumerating a full grid.", [
    "Early trials include anchors and broad exploration.",
    "Later trials are increasingly guided by TPE's model of promising and poor parameter regions.",
    "Approximately 100 × 14 × 3 = 4,200 seed-specific XGBoost fits may be performed.",
    "MLflow records parameters and metrics; Optuna performs the actual selection.",
])
add_step(doc, 19, "Save the lowest-J configuration", "Create a reproducible handoff for fixed analysis.", [
    "Save ESN and XGBoost hyperparameters, preprocessing provenance, seeds and analysis window.",
    "Save the search history, pooled metrics, fold diagnostics and OOF predictions.",
    "Do not overwrite the original non-baseline-centered workflow outputs.",
])

doc.add_page_break()
doc.add_heading("13. Fixed Analysis in AUX-BB", level=1)
add_body(doc, "AUX-BB loads the saved AUX configuration and performs no new hyperparameter search. It reproduces baseline centering, fold-local standardization, ESN feature generation and XGBoost fitting with fixed shared settings.")
add_table(doc, ["Protocol", "Folds", "Test trials per fold", "Question answered"], [
    ("LORO", "6", "56", "Can the model predict a new repetition when materials and temperatures are represented in training?"),
    ("LOMO", "14", "24", "Can the model predict a completely unseen material?"),
    ("LOTO", "4", "84", "Can the model predict at an unseen operating temperature?"),
], [0.8, 0.65, 1.15, 3.9])
add_callout(doc, "Primary scientific protocol", "Use pooled LOMO OOF performance as the main diagnostic of unseen-material generalization.", PALE_TEAL)

doc.add_heading("14. What ‘Unified’ Means", level=1)
add_body(doc, "Unified means that one shared hyperparameter combination is applied to the complete multi-temperature workflow and every material fold. It does not mean that one already-fitted XGBoost object is reused in every fold.")
add_body(doc, "Each fold fits new regressors because its 13-material training set is different. Across 14 LOMO folds and three seeds, 42 fitted regressors collectively produce the 336 pooled OOF predictions.")

doc.add_heading("15. Reservoir-Size Overrides", level=1)
add_body(doc, "A manual override such as OVERRIDE_RES_SIZE = 20 is technically supported because AUX-BB regenerates features and retrains XGBoost. However, the remaining hyperparameters were originally selected jointly with reservoir size 69.")
add_callout(doc, "Reporting rule", "Describe an override as a reservoir-size sensitivity experiment, not as the exact AUX-optimized model. For rigorous fixed-size optimization, rerun AUX with the chosen reservoir size held constant while Optuna tunes the remaining parameters.", PALE_GOLD)

doc.add_page_break()
doc.add_heading("16. Leakage Protections", level=1)
for item in (
    "Material groups do not overlap between training and test portions of a LOMO fold.",
    "Contact detection uses sensor behavior, not k or effusivity.",
    "Per-trial baseline centering uses only that trial's first input sample.",
    "StandardScaler is fitted only on training-material sequences.",
    "XGBoost fits only on training feature rows and training targets.",
    "Held-out target values are used only after prediction for evaluation.",
    "The same preprocessing definition is enforced in both AUX and AUX-BB.",
):
    add_bullet(doc, item)

doc.add_heading("17. Interpretation Limitations", level=1)
add_callout(doc, "Model-selection optimism", "AUX uses the LOMO folds to select hyperparameters, and AUX-BB can reuse the same data and fold structure. These results are cross-validated model-selection performance, not validation on a completely untouched external dataset.", PALE_RED)
add_body(doc, "A stronger final estimate would use nested cross-validation or a separate external dataset that never participates in preprocessing decisions, hyperparameter selection or model development.")
add_body(doc, "Individual reservoir units do not have predetermined physical meanings. Their value is computational: together they provide diverse nonlinear temporal representations. Scientific interpretation should compare input channels, state patterns and prediction errors across materials rather than assigning a physical property to one unit.")

doc.add_heading("18. Final Deployable Ensemble", level=1)
add_body(doc, "After the methodology is selected, a final ensemble may be fitted using all 336 trials. This is distinct from cross-validation.")
deploy_steps = (
    "Baseline-center all trials and fit one final five-channel scaler.",
    "Create the three final reservoirs using seeds 42, 43 and 44.",
    "Generate the three full-data ESN feature tables.",
    "Train three final XGBoost regressors using all known targets.",
    "Save the preprocessing definition, scaler, reservoirs, feature order and regressors.",
    "For a future trial, reproduce contact alignment, 0–5 s extraction and preprocessing; then average the three predictions.",
)
for step in deploy_steps:
    add_number(doc, step)
add_callout(doc, "Performance reporting", "Do not report the final ensemble's training accuracy as expected generalization. Use previously obtained OOF or independent-test performance.")

doc.add_page_break()
doc.add_heading("19. Beginner’s Glossary", level=1)
add_table(doc, ["Term", "Plain-language meaning"], [
    ("Feature", "A numerical input used by a model."),
    ("Target", "The value the model is trained to predict; here, effusivity."),
    ("Regression", "Prediction of a continuous numerical value."),
    ("Fold", "One training/test split within cross-validation."),
    ("LOMO", "Leave one complete material out for testing."),
    ("Out-of-fold prediction", "A prediction made for a trial not used to fit that fold's model."),
    ("Data leakage", "Allowing test information to influence training or preprocessing."),
    ("StandardScaler", "Training-derived mean/standard-deviation transformation."),
    ("Reservoir unit", "One internal ESN state variable evolving through time."),
    ("Seed", "A number controlling reproducible random initialization."),
    ("Hyperparameter", "A model setting chosen outside ordinary fitting."),
    ("Optuna TPE", "A guided search method that proposes promising hyperparameters."),
    ("XGBoost", "A nonlinear ensemble of decision trees used for regression."),
    ("Ensemble", "Multiple fitted models whose predictions are combined."),
    ("OOF", "Out of fold."),
], [1.65, 4.85])

doc.add_page_break()
doc.add_heading("20. Final Methodology Statement", level=1)
add_body(doc, "The complete workflow estimates thermal effusivity from baseline-centered, post-contact thermal transients. Five leakage-safe standardized channels drive three seeded echo-state reservoirs. Compact statistical summaries of the reservoir trajectories, together with operating temperature, are supplied to XGBoost regressors. A shared ESN–XGBoost hyperparameter configuration is selected by Optuna using pooled LOMO out-of-fold accuracy and across-material stability. AUX-BB then applies the fixed configuration to routine LORO, LOMO and LOTO diagnostics, with pooled LOMO performance serving as the principal measure of generalization to unseen materials.")

# Keep tables from splitting individual rows and set metadata.
for table in doc.tables:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)

doc.core_properties.title = "Overall Methodology: AUX ESN–XGBoost Effusivity Regression"
doc.core_properties.subject = "End-to-end beginner-friendly methodology reference"
doc.core_properties.author = "Research workflow documentation"
doc.core_properties.keywords = "ESN, XGBoost, thermal effusivity, LOMO, Optuna, methodology"

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
